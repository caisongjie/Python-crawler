import os
import re
import xlwt
import time
from docx import Document
path = "C:\\Users\\songjie\\Desktop\\qwer"
for files in os.walk(path):
    # print(files[2])
    print(len(files[2]))
used = 0
# for a in range(len(files[2])):\

wr = xlwt.Workbook()
sheet1 = wr.add_sheet('Sheet1', cell_overwrite_ok=True)
for q in range(0,len(files[2])):
    doc = Document("C:\\Users\\songjie\\Desktop\\qwer\\%s"%files[2][q])
    print("processing:%s" % files[2][q])
    print(files[2][q])
    for x in range(1, len(doc.tables[0].rows)):
        print(x)
        print("1 %s" % time.time())
        a = doc.tables[0].rows[x]
        print("2 %s" % time.time())
        x -= 1
        name = a.cells[1].paragraphs[0].text
        sheet1.write(x+used, 0, name)

        temp = re.findall('(\d+)', files[2][q])
        year = temp[0]
        sheet1.write(x+used, 3, year)
        certificate_number = a.cells[2].paragraphs[0].text
        sheet1.write(x+used, 1, certificate_number)
        province = re.findall('(.*?)\d+', files[2][q])
        sheet1.write(x+used, 2, province)
        if re.search('二',files[2][q]) != None:
            batch = 2
        else:
            if re.search('三',files[2][q]) != None:
                batch = 3
            else:
                batch = 1
        sheet1.write(x+used, 4, batch)
    wr.save('C:\\Users\\songjie\\Desktop\\xxx.xls')
    used = used + len(doc.tables[0].rows) - 1